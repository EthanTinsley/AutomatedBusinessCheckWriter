from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from .check import *
from .account import *
from .company import *
from .bank import *
from .database_helper import *
from num2words import num2words


# this class is used in order to manipulate, format, and save the buisness checks
# object can be called and does not require any input information in order to 
# instantiate the check-writer obeject
# check-writer object has methods that accept check parameters in order to manipulate 
# each section of the check template
class CheckWriter:

    def __init__(self):
        pass 

    
    # method for printing a check or checkbook
    # accepts a list of checks to be printed as input parameter
    # accepts an output directory as input parameter
    # foreach check in checklist get signee, account, company, bank from db
    # foreach check in checklist format_check
    # foreach check in checklist print to output directory
    def print_check(self, checklist, outputdir):
        # localize variables
        outputdir = outputdir
        checklist = checklist

        # check template path
        template_path = "./frontend/backend/check_template/CheckTemplate.docx"

        # establish db conenction 
        db = DataBaseHelper()

        # determine if check list to print contains duplicates 
        duplicatelist = self.duplicate_verification(checklist)

        if duplicatelist != [] :
            checklist = self.combine_multicharge_checks(checklist, duplicatelist)
            print(len(checklist))
            print(checklist)

        # iterate through checks in checklist
        for check in checklist:
            # create the current check document
            doc = Document(template_path)

            # determine if check is a multiline check
            if isinstance(check, list):
                # get account information
                account = db.get_account(check[0].account_code)
                # get bank infromation
                bank = db.get_bank(check[0].account_code)
                # get company information
                company = db.get_company(check[0].account_code)
                # get signature
                signature = db.get_signature(check[0].account_code)

                # format the check accordingly
                doc = self.format_multicharge_check(doc, check, company, bank, account, signature)

                # name the check document accordingly
                doc_name = check[0].check_num + '_' + check[0].payee
            
            else:
                # get account information
                account = db.get_account(check.account_code)
                # get bank infromation
                bank = db.get_bank(check.account_code)
                # get company information
                company = db.get_company(check.account_code)
                # get signature
                signature = db.get_signature(check.account_code)

                # format the check accordingly
                doc = self.format_check(doc, check, company, bank, account, signature)

                # name the check document accordingly
                doc_name = check.check_num + '_' + check.payee

                db.update_check_printed(check)
            
            # save the formatted check to the output directroy
            doc.save(outputdir + '/' + doc_name + '.docx')


    # -- edit 8-4-2021 -- 
    # replaced indvidual calls to elif in this call decreased time complexity
    # create a method to alter the CheckTemplate document that will 
    # return a check that has all the corresponding information correctly 
    # formatted on the checking document 
    # returns a docx Document object 
    # Input parameters 
    # the check to be printed 
    # the corresponding company infomration 
    # the corresponding bank infomration 
    # the corresponding account information 
    def format_check(self, check_template_doc, check, company, bank, account, signature):
        # localize input variables
        doc = check_template_doc 
        company = company
        bank = bank
        account = account
        check = check
        signature = signature


        # parse document for keywords to replace 
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # condiitonal statements to locate keywords
                        if "company_name" in para.text:
                            self.format_company_name(para, company.name)
                        elif "address" in para.text:
                            self.format_company_address(para, company.address)
                        elif "city, state zip" in para.text:
                            self.format_company_address2(para, company.city, company.state, company.zip)
                        elif "check_date" in para.text:
                            self.format_check_date(para, check.check_date, 12)
                        elif "date" in para.text:
                            self.format_check_date(para, check.check_date, 8)
                        elif "check_num" in para.text:
                            self.format_check_num(para, check.check_num)
                        elif "payee" in para.text:
                            self.format_payee(para,check.payee)
                        elif "amount_in_text" in para.text:
                            self.format_amount_text(para, check.amount)
                        elif "amount" in para.text:
                            self.format_amount(para, check.amount)
                        elif "Invoice No." in para.text:
                            self.format_memo_table(check, table)
                        elif "bank_name" in para.text:
                            self.format_bank_name(para,bank.name)
                        elif "signature" in para.text:
                            self.format_signature(para, signature.path)
                        elif "routing_num" in para.text:
                            self.format_routing_num(para, account.routing_num)
                        elif "MICR_checknum_BAND" in para.text:
                            self.format_MICR_check_num(para, check.check_num)
                        elif "MICR_accountnum_BAND" in para.text:
                            self.format_MICR_account_num(para, account.account_num)
                        elif "MICR_routingnum_BAND" in para.text:
                            self.format_MICR_routing_num(para, account.routing_num)

        return doc



    # create a method to manipulate a document by adding the comapany information
    # to the check template for printing. In order to do this we will need to accept
    # the check template that needs manipulating along with the company obeject as 
    # input parameters. The method will then parse the input document and add the 
    # comapny information accordingly using some sub-routine methods for a cleaner method
    # the method will then return the altered document after the parsing is complete
    # subroutines - format_company_name() , format_company_address(), format_company_address2()
    def format_company(self, check_template_doc, company):
        doc = check_template_doc 
        name = company.name
        address = company.address
        city = company.city
        state = company.state
        zip = company.zip

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:

                        # parse and replace company information
                        if "company_name" in para.text:
                            self.format_company_name(para, name)
                        elif "address" in para.text:
                            self.format_company_address(para, address)
                        elif "city, state zip" in para.text:
                            self.format_company_address2(para, city, state, zip)
                        elif "bank_name" in para.text:
                            self.format_bank_name(para,name)



        # return the manipulated document 
        return doc


    # subroutine - 1
    # method formats the company name on the document
    # called by format_company()
    def format_company_name(self, para, name):

        # replace placeholder text with the companies name
        para.text = para.text.replace("company_name", name)

        # end method 


    # subroutine - 2
    # method formats the company address on the check 
    # called by format_company()
    def format_company_address(self, para, address):
        
        # remove address placeholder
        para.text = para.text.replace("address", '')

        # create run object to control font size and address name
        run = para.add_run()
        run.font.size = Pt(10)
        run.text = address
        
        # end method 


    # subrountine - 3
    # method formats the companies city , state, and zip code address information
    # clled by format_company()
    def format_company_address2(self, para, city, state, zip):
        # find and replace the city, state zip placeholder
        para.text = para.text.replace("city, state zip" , '')

        # use run object to add in text and font/size
        run = para.add_run()
        run.font.size = Pt(9)
        run.text = city + "," + state + " " + zip

        # end method


    # create a method to format the banks information on the check template 
    # we need to retrieve the banks name for use in printing the check.
    #  The method useds a series of subrotuinesto manipulate the document
    #  for a cleaner code presnetation in this method 
    # subroutines - format_bank_name() 
    def format_bank(self, check_template_doc, bank):
        name = bank.name
        doc = check_template_doc

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        #parse document for the bank_name placeholder
                        if "bank_name" in para.text:
                            self.format_bank_name(para,name)

        # return altered document 
        return doc 

    
    # method to use within format_bank()
    # replaces placeholder text with banking information
    # subroutine - 1 
    def format_bank_name(self, para, name):
        # remove placeholder text
        para.text = ""

        # create run to handle the styling and insertion of banks name
        run = para.add_run()
        run.font.size = Pt(16)
        run.text = name

        # end method

    
    # method to be used when altering invoice and check information within 
    # the CheckTemplate document. uses a series of sub-method / subroutines in order 
    # to properly alter the document to meet required formatting standards
    # returns the altered document after alterations have been made 
    # for input into the method we need the document being altered (CheckTemplate.docx)
    # and the check with all its information Check() object 
    def format_check_info(self, check_template_doc, check):
        doc = check_template_doc
        checkdate = check.check_date
        checknum = check.check_num
        payee = check.payee
        amount = check.amount

        # parse document to add information accordingly
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "check_date" in para.text:
                            self.format_check_date(para, checkdate, 12)
                        elif "date" in para.text:
                            self.format_check_date(para, checkdate, 8)
                        elif "check_num" in para.text:
                            self.format_check_num(para, checknum)
                        elif "payee" in para.text:
                            self.format_payee(para,payee)
                        elif "amount_in_text" in para.text:
                            self.format_amount_text(para, amount)
                        elif "amount" in para.text:
                            self.format_amount(para, amount)
                        elif "Invoice No." in para.text:
                            self.format_memo_table(check, table)
                       # elif "signature" in para.text:
                           # self.format_signature(para, signature.path)
                        
        # return altered document 
        return doc 

    
    # subroutien used in format_check() method 
    # manipulates the document by adding in the check date to the appropriate fields
    # uses run object in order to control formatting 
    # subroutine - 1    
    def format_check_date(self, para, date, fontsize):
        # remove the check_date place holder text 
        para.text = para.text.replace("check_date", '')

        # use a run object to control source formatting 
        run = para.add_run()
        run.font.size = Pt(fontsize)
        run.text = date

        # end method

    # subroutien used in format_check() method 
    # manipulates the document by adding in the check number to the appropriate fields
    # uses run object in order to control formatting 
    # subroutine - 2    
    def format_check_num(self, para, num):
        # remove the check_num placeholder text
        para.text = para.text.replace("check_num", '')

        # use a run object in order to set font style/size and add text 
        run = para.add_run()
        run.text = "Check No. " + num

        # end method 


    # subroutien used in format_check() method 
    # manipulates the document by adding in the payee to the appropriate fields
    # uses run object in order to control formatting 
    # subroutine - 3    
    def format_payee(self, para, payee):
        # remove the payee placholder text 
        para.text = para.text.replace("payee", '')

        # use run object in order to format the text 
        run = para.add_run()
        run.text = payee

        # end method 


    # subroutine used in format_check() method 
    # manipulates the document by adding in the amount of the check in the appropriate places 
    # uses run objct in order to control formtting 
    # subroutine - 4
    def format_amount(self, para, amount):
        # remove the amount place holder 
        para.text = para.text.replace("amount", '')

        # create a string to place in the document
        amounttext = "" + str(amount)

        # use run object to control the formatting of the text 
        run = para.add_run()
        run.font.size = Pt(8)
        run.text = amounttext

        # end the method here        

    # subrountine used in format_check() method
    # mainpulates the document in oder to add the amount of the check in text
    # uses num2words library in order to convert the amount of the check into a string
    # of words that are equal to the amount (float) of the check
    # formats the cents of the check as a fraction out of 100 to acheive this 
    # we subtract the amoutn of the check (float) by the amount of the check (int) 
    # in order to allienate the decimal and multiply it by 100 to make it a whole number
    # additionally, we need to round this number to 2 decimal places as using 01 will result in 
    # in an outcome of 0 and not 1.
    def format_amount_text(self, para, amount):
        # localize variables 
        para = para
        amount = amount 

        # convert the amount of the check to words with num2words library
        amounttext = num2words(int(float(amount)))
        amounttext = amounttext + " DOLLARS & "

        # find the amount of cents and make it into a whole number (multiply by 100 and round to 2 decimal places)
        decimal = int(round(((float(amount) - int(float(amount)))* 100),2))
        decimaltext = str(decimal) + "/100"

        # combine the dollar and cent amount into one string
        amounttext = "PAY " + amounttext + decimaltext
        amounttext = amounttext.upper()

        # remove palceholder text from the document
        para.text = para.text.replace("amount_in_text", '')

        # add the amount text to the document using a run object
        run = para.add_run()
        #run.font.size = Pt(12)
        run.text = amounttext

        #end method


    # a method used to format the accounting information on the check
    # will parse throug the document to find neccesary information and using sub-methods
    # will replace the placeholder text with the neccesary information 
    # will change routing number under bank name 
    # will change all of MICR band 
    def format_account(self, check_template_doc, account, check_num):
        # localize variables
        account = account 
        doc = check_template_doc
        check_num = check_num

        # parse the document to find key terms (IE location)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # check for accounting information within each paragraph
                        if "routing_num" in para.text:
                            self.format_routing_num(para, account.routing_num)
                        elif "MICR_checknum_BAND" in para.text:
                            self.format_MICR_check_num(para, check_num)
                        elif "MICR_accountnum_BAND" in para.text:
                                self.format_MICR_account_num(para, account.account_num)
                        elif "MICR_routingnum_BAND" in para.text:
                            self.format_MICR_routing_num(para, account.routing_num)

        # end of method
        return doc


    # method will format the routing number underneath the banking logo 
    # addiotnally, will insert "ACH R/T" for desingation as to what number is
    # is a sub-method used within format_account()
    def format_routing_num(self, para, routing_num):
        num = routing_num

        # format routing number as a string and add acknowledgement to it
        routing = "ACH R/T " + str(num)

        # remove placeholder
        para.text = para.text.replace("routing_num", '')

        # add text into paragrapgh using run object 
        run = para.add_run()
        run.font.size = Pt(8)
        run.text = routing

    
    # method that will add the appropriate signature image to the check in the signature box
    def format_signature(self, para, signature_img):
        para = para
        img_path = signature_img

        # remove plaveholder text in paragraph
        para.text = para.text.replace('signature', '')

        # create run object and place in paragraph
        run = para.add_run()
        run.add_picture(img_path, height=Inches(.50))

        # end method


    # create a method that will format the checknumber on the MICR band of the check 
    # method will accept the check number from the format_account() method
    # method will parsethe string to retrieve each of the digits of the routing number and 
    # for each digit in the string will insert an MICR charachter via an image path method
    # can set a minimum check number and wil place leading 0's to bring the check number up
    # to the supported minimum length
    def format_MICR_check_num(self, para, check_num):
        num = str(check_num)

        # set a minimum length to the check number MICR band
        min_len = 6
        
        # determine the number of leading 0's to use
        zeros = min_len - len(num)

        # add the leading 0's to the check number
        if zeros > 0:
            for x in range(zeros):
                num = '0' + num

        # remove placeholder text 
        para.text = para.text.replace("MICR_checknum_BAND", '')

        # intialize run object in order to insert charachters into the document 
        run = para.add_run()
        run.font.size = Pt(8)

        # insert leading 'on-us' symbol
        path = self.get_MICR_char_path('on-us')
        run.add_picture(path, height=Inches(.12))
        run.add_text(' ')

        # use for-loop to insert each image/digit into the document 
        # use the get_MICR method to retreive the path of the appliacable digits 
        for char in num:
            
            path = self.get_MICR_char_path(char)
            run.add_picture(path, height=Inches(.125))
            run.add_text(' ')

        # insert trailing 'on-us' symbol
        path = self.get_MICR_char_path('on-us')
        run.add_picture(path, height=Inches(.12))
        run.add_text(' ')
        
        # end method


    # method that formats the routing number on the check using MICR charachters
    # accepts the paragrpagh to insert into and the routing number as input parameters 
    # uses a for loop to place each digit of the account number into the document 
    # using a docx run object
    # sub-method used by format_account()
    def format_MICR_routing_num(self, para, routing_num):
        num = str(routing_num)

        # set a minimum length to the check number MICR band
        min_len = 6
        
        # determine the number of leading 0's to use
        zeros = min_len - len(num)

        # add the leading 0's to the check number
        if zeros > 0:
            for x in range(zeros):
                num = '0' + num

        # remove placehodler text
        para.text = para.text.replace('MICR_routingnum_BAND', '')

        # create a run object and place the transit MICR charachter 
        run = para.add_run()
        run.font.size = Pt(10)

        # path to the transit charachter 
        path = self.get_MICR_char_path('transit')
        run.add_picture(path, height= Inches(.125))
        run.add_text(' ')

        # insert the charachters into the paragraph one by one with a for-each loop
        for char in num:

            path = self.get_MICR_char_path(char)
            run.add_picture(path, height=Inches(.125))
            run.add_text(' ')

        # add closing transit symbol
        # path to the transit charachter 
        path = self.get_MICR_char_path('transit')
        run.add_picture(path, height= Inches(.125))

        # end method


    # method that will add the account number to the MICR line
    # accepts the paragrapgh that will be manipulated 
    # accepts the account number being added 
    # runs the account number through a for-each loop that replaces
    # each of the digits with their MICR image using the MICR_lookup method
    # used within format_account()
    def format_MICR_account_num(self, para, account_num):
        num = str(account_num)

        # set a minimum length to the check number MICR band
        min_len = 6
        
        # determine the number of leading 0's to use
        zeros = min_len - len(num)

        # add the leading 0's to the check number
        if zeros > 0:
            for x in range(zeros):
                num = '0' + num

        # remove placehodler text
        para.text = para.text.replace('MICR_accountnum_BAND', '')

        # create a run object and place the transit MICR charachter 
        run = para.add_run()
        run.font.size = Pt(10)

        # insert the charachters into the paragraph one by one with a for-each loop
        for char in num:

            path = self.get_MICR_char_path(char)
            run.add_picture(path, height=Inches(.125))
            run.add_text(' ')

        # add closing transit symbol
        # path to the transit charachter 
        path = self.get_MICR_char_path('transit')
        run.add_picture(path, height= Inches(.125))

        # end method


    # method used to return the path for the desired MICR E-13B font charachter     
    # if input number is '0' will return the path to the MICR '0' image 
    def get_MICR_char_path(self, char):
        char = char 
        path = './frontend/backend/micr_char/'

        if char == '0':
            path += 'MICR_0.png'
        elif char == '1':
            path += 'MICR_1.png'
        elif char == '2':
            path += 'MICR_2.png'
        elif char == '3':
            path += 'MICR_3.png'
        elif char == '4':
            path += 'MICR_4.png'
        elif char == '5':
            path += 'MICR_5.png'
        elif char == '6':
            path += 'MICR_6.png'
        elif char == '7':
            path += 'MICR_7.png'
        elif char == '8':
            path += 'MICR_8.png'
        elif char == '9':
            path += 'MICR_9.png'
        elif char == 'on-us':
            path += 'MICR_on-us.png'
        elif char == 'transit':
            path += 'MICR_transit.png'

        return path

    
    # method to format the checks memo table for a singular check
    # should take the information from the check object and insert it into the 
    # table that will house the memo information of the check 
    # accepts the document to manipulate as well as the check object to be input 
    # used in the format_check() method
    def format_memo_table(self, check, table):
        check = check
        count = 0
        memotable = table
        memo_row = table.rows[1]
                        
        # iterate through the memo table and add the check information to the memo 
        # add key information to the memo table depending on the column count 
        # of current row
        row = memo_row
        for cell in row.cells:
            para = cell.paragraphs[0]

            if count == 0: # cell for invoice number
                run = para.add_run()
                run.font.size = Pt(10)
                run.add_text(str(check.invoice_num))
            elif count == 1: # cell for invoice date 
                run = para.add_run()
                run.font.size = Pt(10)
                run.add_text(str(check.invoice_date))
            elif count == 2: # cell for reference 
                para.text = ''
            elif count == 3: # cell for account information
                para.text = ''
            elif count == 4: # cell for gross amount 
                run = para.add_run()
                run.font.size = Pt(10)
                run.add_text(str(check.amount))
            elif count == 5: # cell for discount 
                para.text = ''
            elif count == 6: # cell for total balance 
                run = para.add_run()
                run.font.size = Pt(10)
                run.add_text(str(check.amount))
            
            count += 1 # move to the next cell

        # end of the method



    # a method used to format a multi-charge check
    # accepts the same parameters as format_check() accept that 
    # instead of a check object we are accepting a list of check objects
    def format_multicharge_check(self, check_template_doc, check_list, company, bank, account, signature): 
        # localize variables
        doc = check_template_doc 
        company = company
        bank = bank
        account = account
        checklist = check_list
        signature = signature
        amount = 0

        for check in checklist:
            amount += float(check.amount)
        
        # parse document for keywords to replace 
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # condiitonal statements to locate keywords
                        if "company_name" in para.text:
                            self.format_company_name(para, company.name)
                        elif "address" in para.text:
                            self.format_company_address(para, company.address)
                        elif "city, state zip" in para.text:
                            self.format_company_address2(para, company.city, company.state, company.zip)
                        elif "check_date" in para.text:
                            self.format_check_date(para, checklist[0].check_date, 12)
                        elif "date" in para.text:
                            self.format_check_date(para, checklist[0].check_date, 8)
                        elif "check_num" in para.text:
                            self.format_check_num(para, checklist[0].check_num)
                        elif "payee" in para.text:
                            self.format_payee(para,checklist[0].payee)
                        elif "amount_in_text" in para.text:
                            self.format_amount_text(para, amount)
                        elif "amount" in para.text:
                            self.format_amount(para, amount)
                        elif "Invoice No." in para.text:
                            self.format_multiline_memo_table(checklist, table)
                        elif "bank_name" in para.text:
                            self.format_bank_name(para,bank.name)
                        elif "signature" in para.text:
                            self.format_signature(para, signature.path)
                        elif "routing_num" in para.text:
                            self.format_routing_num(para, account.routing_num)
                        elif "MICR_checknum_BAND" in para.text:
                            self.format_MICR_check_num(para, checklist[0].check_num)
                        elif "MICR_accountnum_BAND" in para.text:
                            self.format_MICR_account_num(para, account.account_num)
                        elif "MICR_routingnum_BAND" in para.text:
                            self.format_MICR_routing_num(para, account.routing_num)

        return doc

    
    # method to format the checks memo table for a multi-line check
    # should take the information from the check object and insert it into the 
    # table that will house the memo information of the check 
    # accepts the document to manipulate as well as the check object to be input 
    # the desired row is where the checks information will be inserted 
    # the total is the what will be tallied in the net column of the table
    # used in the format_check() method
    def format_multiline_memo_table(self, checklist, table):
        memotable = table
        total = 0
        count = 0
        index = -1
                        
        # iterate through the memo table and add the check information to the memo 
        # add key information to the memo table depending on the column count 
        # of current row
        for row in memotable.rows:
            if index == -1:
                index += 1
                continue
            elif index == len(checklist):
                break
            
            total += checklist[index].amount

            for cell in row.cells:
                para = cell.paragraphs[0]

                if count == 0: # cell for invoice number
                    run = para.add_run()
                    run.font.size = Pt(10)
                    run.add_text(str(checklist[index].invoice_num))
                elif count == 1: # cell for invoice date 
                    run = para.add_run()
                    run.font.size = Pt(10)
                    run.add_text(str(checklist[index].invoice_date))
                elif count == 2: # cell for reference 
                    para.text = ''
                elif count == 3: # cell for account information
                    para.text = ''
                elif count == 4: # cell for gross amount 
                    run = para.add_run()
                    run.font.size = Pt(10)
                    run.add_text(str(checklist[index].amount))
                elif count == 5: # cell for discount 
                    para.text = ''
                elif count == 6: # cell for total balance 
                    run = para.add_run()
                    run.font.size = Pt(10)
                    run.add_text(str(round(float(total), 2)))
            
                count += 1 # move to the next cell

            count = 0
            index += 1 # move to the next check


        # end of the method
                

    # method for determing if list of checks contains check_num duplicates 
    # a duplicate in check_num would illustrate that a multi-line charge check is present in the list 
    # parse through checklist for print and determine if duplicates exist
    # if duplicate found then combine to a duplicate list and return the list of duplicates
    def duplicate_verification(self, checklist):
        
        dupchecknums = []
        set_checknums = set()
        for check in checklist:
            if check.check_num in set_checknums:
                dupchecknums.append(check.check_num)
            else:
                set_checknums.add(check.check_num)
        
        return dupchecknums

    
    # method accepts list of duplicate check numbers to determine whihc checks to combine 
    # method accepts list of checks and combines the mutliline charge checks 
    def combine_multicharge_checks(self, checklist, dupnumlist):
        templist = checklist[:]

        for duplicate in dupnumlist:
            currlist = []
            for check in checklist:
                if check.check_num == duplicate:
                    currlist.append(check)
                    templist.remove(check)

            templist.append(currlist)

        # retubr modified check list
        return templist
